from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


@dataclass
class PaperConfig:
    department: str
    subject: str
    subject_code: str
    semester: str
    max_marks: int
    duration: str
    date: str
    batch: str
    teaching_department: str
    exam_type: str
    modules: list[int]
    rbt_levels: list[str]
    co_targets: list[str]
    year: str = "2026"
    instructions: str = "Instruction: Answer the following questions"
    college_name: str = "Dayananda Sagar Academy of Technology & Management"
    affiliation: str = "(Autonomous Institute under VTU)"
    program_line: str = "6 Programs Accredited by NBA (CSE, ISE, ECE, EEE, MECH, CV)"
    co_descriptions: dict[str, str] = field(default_factory=dict)
    co_percentages: dict[str, int] = field(default_factory=dict)
    module_percentages: dict[str, int] = field(default_factory=dict)
    left_seal_label: str = "DSATM"
    right_seal_label: str = "IQAC"
    template_note: str | None = None
    template_family: str = "dsatm"


import random
from .academic.policies import get_allowed_rbt

QUESTION_PATTERNS = [
    {
        "type": "3_sub",
        "distribution": [6, 6, 8]
    },
    {
        "type": "3_sub",
        "distribution": [8, 8, 4]
    },
    {
        "type": "2_sub",
        "distribution": [10, 10]
    },
    {
        "type": "single",
        "distribution": [20]
    }
]

STRUCTURE_WEIGHTS = {
    "single": 0.4,
    "2_sub": 0.4,
    "3_sub": 0.2
}

def build_question_blueprint(max_marks: int, module_co_map: dict[int, str] | None = None, module_image_map: dict[int, bool] | None = None) -> list[dict[str, Any]]:
    """
    Dynamically generates a randomized question paper blueprint.
    Ensures exactly ONE module uses a 3-sub-question format, while others
    use a combination of single and 2-sub based on probabilities.
    """
    blueprint: list[dict[str, Any]] = []
    
    # 1. Determine which module gets the 3_sub format
    modules = list(range(1, 6))
    three_sub_module = random.choice(modules)
    
    question_number = 1
    for module_number in modules:
        if module_co_map and module_number in module_co_map:
            co_target = module_co_map[module_number]
        else:
            co_target = f"CO{module_number}"
        allowed_rbt = get_allowed_rbt(co_target)
        
        # 2. Select structure for this module based on constraints
        if module_number == three_sub_module:
            # Must be 3_sub
            available_patterns = [p for p in QUESTION_PATTERNS if p["type"] == "3_sub"]
            chosen_pattern = random.choice(available_patterns)
            chosen_distribution = chosen_pattern["distribution"]
        else:
            # Must NOT be 3_sub
            available_patterns = [p for p in QUESTION_PATTERNS if p["type"] != "3_sub"]
            weights = [STRUCTURE_WEIGHTS[p["type"]] for p in available_patterns]
            chosen_pattern = random.choices(available_patterns, weights=weights, k=1)[0]
            chosen_distribution = chosen_pattern["distribution"]
            
        # Image question flag
        wants_image = module_image_map and module_image_map.get(module_number, False)
        # Assign the image to only one subpart randomly if requested
        image_assigned = False

        # 3. Apply the structure twice (internal choice)
        for q_index in range(2):
            subpart_labels = ["a", "b", "c", "d"]
            for i, marks in enumerate(chosen_distribution):
                subpart = subpart_labels[i] if len(chosen_distribution) > 1 else ""
                rbt = random.choice(allowed_rbt)
                
                is_img_q = False
                if wants_image and not image_assigned:
                    # roughly 50% chance to assign it to this slot until we run out
                    if random.random() > 0.5 or (q_index == 1 and i == len(chosen_distribution) - 1):
                        is_img_q = True
                        image_assigned = True
                
                blueprint.append(
                    {
                        "question_number": question_number,
                        "subpart": subpart,
                        "label": format_question_label(question_number, subpart) if subpart else str(question_number),
                        "marks": marks,
                        "module_number": module_number,
                        "co": co_target,
                        "rbt": rbt,
                        "is_image_question": is_img_q,
                    }
                )
            question_number += 1
            
    return blueprint


def format_question_label(question_number: int, subpart: str) -> str:
    return f"{question_number}({subpart})"


def normalize_question_label(label: Any) -> str:
    text = str(label or "").strip()
    if len(text) >= 2 and text[-1:].isalpha() and text[:-1].isdigit():
        return format_question_label(int(text[:-1]), text[-1:])
    return text


def build_question_rows(
    max_marks: int, questions: list[dict[str, Any]]
) -> list[dict[str, Any]]:
    blueprint = build_question_blueprint(max_marks)
    padded_questions = questions[: len(blueprint)] + [
        {} for _ in range(max(0, len(blueprint) - len(questions[: len(blueprint)])))
    ]
    rows: list[dict[str, Any]] = []
    current_module: int | None = None

    for slot, question in zip(blueprint, padded_questions):
        if max_marks > 50 and slot["module_number"] != current_module:
            current_module = int(slot["module_number"])
            rows.append({"type": "module", "title": f"Module - {current_module}"})

        if (slot["subpart"] == "a" or slot["subpart"] == "") and slot["question_number"] % 2 == 0:
            rows.append({"type": "or"})

        rows.append(
            {
                "type": "question",
                "qno": normalize_question_label(
                    question.get("section_label") or slot["label"]
                ),
                "text": str(question.get("text", "")),
                "marks": int(question.get("marks", slot["marks"]) or slot["marks"]),
                "co": str(question.get("course_outcome", "")),
                "rbtl": str(question.get("bloom_level", "")),
                "figure_image_paths": question.get("figure_image_paths", []),
            }
        )

    return rows


class DSATMQuestionPaperGenerator:
    def generate(
        self, config: PaperConfig, questions: list[dict[str, Any]]
    ) -> DocumentType:
        document = Document()
        self._set_page_layout(document)
        self._add_header(document, config)
        self._add_usn_row(document)
        self._add_department_heading(document, config)
        self._add_exam_title(document, config)
        self._add_meta_table(document, config)
        self._add_instruction(document, config)
        self._add_questions_table(document, config, questions)
        self._add_course_outcomes_table(document, config)
        return document

    def save(self, document: DocumentType, output_path: Path) -> Path:
        output_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(str(output_path))
        return output_path

    def _set_page_layout(self, document: DocumentType) -> None:
        section = document.sections[0]
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Inches(0.35)
        section.bottom_margin = Inches(0.45)
        section.left_margin = Inches(0.4)
        section.right_margin = Inches(0.4)
        section.start_type = WD_SECTION.NEW_PAGE

        normal = document.styles["Normal"]
        normal.font.name = "Arial"
        normal.font.size = Pt(9)
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

    def _style_run(
        self,
        run,
        *,
        size: int = 9,
        bold: bool = False,
        italic: bool = False,
        color: RGBColor | None = None,
    ) -> None:
        run.font.name = "Arial"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        if color is not None:
            run.font.color.rgb = color

    def _set_table_borders(
        self,
        table,
        *,
        top: bool = True,
        left: bool = True,
        bottom: bool = True,
        right: bool = True,
        inside_h: bool = True,
        inside_v: bool = True,
        size: str = "10",
    ) -> None:
        borders = OxmlElement("w:tblBorders")
        mapping = {
            "top": top,
            "left": left,
            "bottom": bottom,
            "right": right,
            "insideH": inside_h,
            "insideV": inside_v,
        }
        for name, enabled in mapping.items():
            border = OxmlElement(f"w:{name}")
            border.set(qn("w:val"), "single" if enabled else "nil")
            border.set(qn("w:sz"), size)
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), "000000")
            borders.append(border)

        table_element = table._tbl
        table_pr = table_element.tblPr
        existing = table_pr.first_child_found_in("w:tblBorders")
        if existing is not None:
            table_pr.remove(existing)
        table_pr.append(borders)

    def _set_table_fixed_layout(self, table) -> None:
        table_pr = table._tbl.tblPr
        existing = table_pr.first_child_found_in("w:tblLayout")
        if existing is not None:
            table_pr.remove(existing)
        layout = OxmlElement("w:tblLayout")
        layout.set(qn("w:type"), "fixed")
        table_pr.append(layout)

    def _set_cell_margins(
        self,
        cell,
        *,
        top: int = 45,
        left: int = 65,
        bottom: int = 45,
        right: int = 65,
    ) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        margins = tc_pr.first_child_found_in("w:tcMar")
        if margins is None:
            margins = OxmlElement("w:tcMar")
            tc_pr.append(margins)
        for side, value in {
            "top": top,
            "start": left,
            "bottom": bottom,
            "end": right,
        }.items():
            node = margins.find(qn(f"w:{side}"))
            if node is None:
                node = OxmlElement(f"w:{side}")
                margins.append(node)
            node.set(qn("w:w"), str(value))
            node.set(qn("w:type"), "dxa")

    def _set_cell(
        self,
        cell,
        text: str,
        *,
        bold: bool = False,
        italic: bool = False,
        size: int = 9,
        align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
        vertical_align: WD_ALIGN_VERTICAL = WD_ALIGN_VERTICAL.CENTER,
    ) -> None:
        cell.text = ""
        self._set_cell_margins(cell)
        lines = str(text).splitlines() or [""]
        for index, line in enumerate(lines):
            paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
            paragraph.alignment = align
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            run = paragraph.add_run(line)
            self._style_run(run, size=size, bold=bold, italic=italic)
        cell.vertical_alignment = vertical_align

    def _add_header(self, document: DocumentType, config: PaperConfig) -> None:
        table = document.add_table(rows=1, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        self._set_table_borders(
            table,
            top=False,
            left=False,
            bottom=False,
            right=False,
            inside_h=False,
            inside_v=False,
        )
        self._set_table_fixed_layout(table)

        widths = [0.8, 3.2, 2.5, 0.8]
        for index, width in enumerate(widths):
            table.columns[index].width = Inches(width)

        left_cell, title_cell, approval_cell, right_cell = table.rows[0].cells
        
        # Load and render left DSATM seal image
        left_cell.text = ""
        p_left = left_cell.paragraphs[0]
        p_left.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_left = p_left.add_run()
        left_seal_path = Path(__file__).parent / "assets" / "dsatm-seal.png"
        if left_seal_path.exists():
            run_left.add_picture(str(left_seal_path), width=Inches(0.72))
        else:
            self._set_seal(left_cell, config.left_seal_label)
        left_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Title Cell
        title_cell.text = ""
        p_title = title_cell.paragraphs[0]
        p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_title.paragraph_format.space_before = Pt(4)
        p_title.paragraph_format.space_after = Pt(4)
        run_name = p_title.add_run("Dayananda Sagar Academy of Technology & Management")
        self._style_run(run_name, size=11, bold=True)
        p_title.add_run("\n")
        run_aff = p_title.add_run("(Autonomous Institute under VTU)")
        self._style_run(run_aff, size=8)
        title_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Approval Cell
        approval_cell.text = ""
        tc_pr = approval_cell._tc.get_or_add_tcPr()
        borders = OxmlElement("w:tcBorders")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "8")  # 1pt border
        left.set(qn("w:color"), "A0A0A0")
        borders.append(left)
        tc_pr.append(borders)
        
        self._set_cell_margins(approval_cell, left=150, top=45, bottom=45, right=45)

        p_app = approval_cell.paragraphs[0]
        p_app.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_app.paragraph_format.space_before = Pt(0)
        p_app.paragraph_format.space_after = Pt(0)
        self._append_highlighted_line(
            p_app,
            [("Affiliated to ", False), ("VTU", True)],
            size=7.5,
        )
        self._append_highlighted_line(
            p_app,
            [("Approved by ", False), ("AICTE", True)],
            size=7.5,
        )
        self._append_highlighted_line(
            p_app,
            [("Accredited by ", False), ("NAAC", True), (" with A+ Grade", True)],
            size=7.5,
        )
        self._append_highlighted_line(
            p_app,
            [("6 Programs Accredited by ", False), ("NBA", True)],
            size=7.5,
        )
        run_parentheses = p_app.add_run("(CSE, ISE, ECE, EEE, MECH, CV)")
        self._style_run(run_parentheses, size=7.2)
        approval_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Load and render right IQAC seal image
        right_cell.text = ""
        p_right = right_cell.paragraphs[0]
        p_right.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_right = p_right.add_run()
        right_seal_path = Path(__file__).parent / "assets" / "iqac-seal.png"
        if right_seal_path.exists():
            run_right.add_picture(str(right_seal_path), width=Inches(0.72))
        else:
            self._set_seal(right_cell, config.right_seal_label)
        right_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Divider line: Solid bottom border across all 4 header cells
        for cell in table.rows[0].cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_borders = tc_pr.first_child_found_in("w:tcBorders")
            if tc_borders is None:
                tc_borders = OxmlElement("w:tcBorders")
                tc_pr.append(tc_borders)
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "12")  # 1.5pt solid divider line
            bottom.set(qn("w:color"), "000000")
            tc_borders.append(bottom)

    def _set_seal(self, cell, label: str) -> None:
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        top = paragraph.add_run(label)
        self._style_run(top, size=10, bold=True)
        bottom = paragraph.add_run("\nSeal")
        self._style_run(bottom, size=7)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    def _append_highlighted_line(
        self,
        paragraph,
        parts: list[tuple[str, bool]],
        *,
        size: float,
    ) -> None:
        for index, (text, emphasized) in enumerate(parts):
            run = paragraph.add_run(text)
            color = RGBColor(198, 40, 40) if emphasized else None
            self._style_run(run, size=int(size), color=color)
            if index == len(parts) - 1:
                paragraph.add_run("\n")

    def _add_usn_row(self, document: DocumentType) -> None:
        table = document.add_table(rows=1, cols=11)
        table.alignment = WD_TABLE_ALIGNMENT.RIGHT
        table.autofit = False
        self._set_table_borders(table, top=False, left=False, bottom=False, right=False)
        self._set_table_fixed_layout(table)

        label_cell = table.rows[0].cells[0]
        label_cell.width = Inches(0.55)
        self._set_cell(
            label_cell,
            "USN:",
            size=8,
            align=WD_ALIGN_PARAGRAPH.RIGHT,
        )
        for index in range(1, 11):
            table.columns[index].width = Inches(0.31)
            self._set_cell(
                table.rows[0].cells[index],
                "",
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
            self._set_table_borders(
                table,
                top=False,
                left=False,
                bottom=False,
                right=False,
                inside_h=False,
                inside_v=False,
            )
            cell = table.rows[0].cells[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            borders = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right"):
                border = OxmlElement(f"w:{side}")
                border.set(qn("w:val"), "single")
                border.set(qn("w:sz"), "8")
                border.set(qn("w:color"), "000000")
                borders.append(border)
            tc_pr.append(borders)

    def _add_department_heading(self, document: DocumentType, config: PaperConfig) -> None:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(8)
        paragraph.paragraph_format.space_after = Pt(8)
        run = paragraph.add_run(f"Department of {config.department}")
        self._style_run(run, size=12, bold=True)

    def _add_exam_title(self, document: DocumentType, config: PaperConfig) -> None:
        table = document.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        self._set_table_borders(table)
        self._set_table_fixed_layout(table)
        table.columns[0].width = Inches(7.35)
        self._set_cell(
            table.rows[0].cells[0],
            config.exam_type,
            bold=True,
            size=10,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )

    def _add_meta_table(self, document: DocumentType, config: PaperConfig) -> None:
        table = document.add_table(rows=2, cols=6)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        self._set_table_borders(table)
        self._set_table_fixed_layout(table)
        widths = [1.1, 2.4, 0.8, 1.0, 1.0, 1.05]
        for index, width in enumerate(widths):
            table.columns[index].width = Inches(width)

        headers = [
            "Course Code:",
            "Course Title:",
            "Semester:",
            "Date:",
            "Max. Marks:",
            "Duration:",
        ]
        values = [
            config.subject_code,
            config.subject,
            config.semester,
            config.date,
            str(config.max_marks),
            config.duration,
        ]

        for i in range(6):
            self._set_cell(table.rows[0].cells[i], headers[i], bold=True, size=8.5, align=WD_ALIGN_PARAGRAPH.LEFT)
            self._set_cell(table.rows[1].cells[i], values[i], bold=False, size=8.5, align=WD_ALIGN_PARAGRAPH.LEFT)

    def _add_instruction(self, document: DocumentType, config: PaperConfig) -> None:
        # Note to Students title
        p = document.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        run_title = p.add_run("Note to Students:")
        self._style_run(run_title, size=9.5, bold=True)

        notes = [
            "1. Answer any five full questions, choosing at least one full question from each part.",
            "2. Use of non-programmable calculators is permitted.",
            "3. Assume any missing data suitably, if any.",
        ]
        for note in notes:
            p_note = document.add_paragraph()
            p_note.paragraph_format.space_before = Pt(0)
            p_note.paragraph_format.space_after = Pt(0)
            p_note.paragraph_format.left_indent = Inches(0.2)
            run_note = p_note.add_run(note)
            self._style_run(run_note, size=9)

        # Thin divider below Note to Students
        divider = document.add_table(rows=1, cols=1)
        divider.alignment = WD_TABLE_ALIGNMENT.CENTER
        divider.autofit = False
        self._set_table_borders(
            divider,
            top=False,
            left=False,
            right=False,
            bottom=True,
            inside_h=False,
            inside_v=False,
            size="8",
        )
        divider.columns[0].width = Inches(7.35)
        divider.rows[0].cells[0].text = ""

    def _add_questions_table(
        self,
        document: DocumentType,
        config: PaperConfig,
        questions: list[dict[str, Any]],
    ) -> None:
        paper_rows = build_question_rows(config.max_marks, questions)
        table = document.add_table(rows=1 + len(paper_rows), cols=5)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        self._set_table_borders(table)
        self._set_table_fixed_layout(table)
        widths = [0.6, 5.0, 0.8, 0.6, 0.6]
        for index, width in enumerate(widths):
            table.columns[index].width = Inches(width)

        headers = ["Q\nNo", "Questions", "Marks", "COs", "RBTL"]
        for index, header in enumerate(headers):
            self._set_cell(
                table.rows[0].cells[index],
                header,
                bold=True,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )

        row_index = 1
        for item in paper_rows:
            if item["type"] == "module":
                module_row = table.rows[row_index]
                merged = module_row.cells[0].merge(module_row.cells[4])
                self._set_cell(
                    merged,
                    item["title"],
                    bold=True,
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                )
                row_index += 1
                continue

            if item["type"] == "or":
                or_row = table.rows[row_index]
                merged = or_row.cells[0].merge(or_row.cells[4])
                self._set_cell(
                    merged,
                    "OR",
                    bold=True,
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                )
                row_index += 1
                continue

            current = table.rows[row_index]
            for cell_index, width in enumerate(widths):
                current.cells[cell_index].width = Inches(width)
            self._set_cell(
                current.cells[0],
                item["qno"],
                align=WD_ALIGN_PARAGRAPH.CENTER,
                vertical_align=WD_ALIGN_VERTICAL.TOP,
            )
            self._set_cell(current.cells[1], item["text"], vertical_align=WD_ALIGN_VERTICAL.TOP)
            if item.get("figure_image_paths"):
                for img_path in item["figure_image_paths"]:
                    if img_path and Path(img_path).exists():
                        p = current.cells[1].add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p.paragraph_format.space_before = Pt(6)
                        p.paragraph_format.space_after = Pt(6)
                        run = p.add_run()
                        run.add_picture(img_path, width=Inches(3.5))
            self._set_cell(
                current.cells[2],
                str(item["marks"]),
                align=WD_ALIGN_PARAGRAPH.CENTER,
                vertical_align=WD_ALIGN_VERTICAL.TOP,
            )
            self._set_cell(
                current.cells[3],
                item["co"],
                align=WD_ALIGN_PARAGRAPH.CENTER,
                vertical_align=WD_ALIGN_VERTICAL.TOP,
            )
            self._set_cell(
                current.cells[4],
                item["rbtl"],
                align=WD_ALIGN_PARAGRAPH.CENTER,
                vertical_align=WD_ALIGN_VERTICAL.TOP,
            )
            current.height_rule = WD_ROW_HEIGHT_RULE.AUTO
            row_index += 1

    def _add_course_outcomes_table(
        self,
        document: DocumentType,
        config: PaperConfig,
    ) -> None:
        heading = document.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading.paragraph_format.space_before = Pt(10)
        heading.paragraph_format.space_after = Pt(2)
        label = heading.add_run(
            "Course Outcomes (COs):  At the end of the Course, the Student will be able to:"
        )
        self._style_run(label, size=8.5, bold=True)

        table = document.add_table(rows=5, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        self._set_table_borders(table)
        self._set_table_fixed_layout(table)
        table.columns[0].width = Inches(0.6)
        table.columns[1].width = Inches(6.7)
        for index in range(1, 6):
            co_key = f"CO{index}"
            self._set_cell(
                table.rows[index - 1].cells[0],
                co_key,
                bold=True,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
            self._set_cell(
                table.rows[index - 1].cells[1],
                config.co_descriptions.get(co_key, ""),
            )


def generate_question_paper(
    config: PaperConfig, questions: list[dict[str, Any]], output_dir: Path
) -> Path:
    generator = DSATMQuestionPaperGenerator()
    document = generator.generate(config, questions)
    filename = (
        f"QP_{config.subject_code}_{config.exam_type}_{config.date.replace('-', '')}.docx"
    )
    return generator.save(document, output_dir / filename)


docx_generator = DSATMQuestionPaperGenerator()
