from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def build_footer(document: Document, course_outcomes: dict):
    """
    course_outcomes: { "CO1": "description", "CO2": "description", ... }
    """
    document.add_paragraph()

    label_para = document.add_paragraph()
    label_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    lr = label_para.add_run("Course Outcomes (COs):")
    lr.bold = True
    lr.font.size = Pt(11)

    co_table = document.add_table(rows=len(course_outcomes), cols=2)
    co_table.style = "Table Grid"

    for idx, (co_key, co_text) in enumerate(course_outcomes.items()):
        row = co_table.rows[idx]

        # CO label cell
        label_cell = row.cells[0]
        label_cell.paragraphs[0].clear()
        p = label_cell.paragraphs[0]
        r = p.add_run(co_key)
        r.bold = True
        r.font.size = Pt(10)
        label_cell.width = 700000   # ~0.5 inch

        # CO description cell
        desc_cell = row.cells[1]
        desc_cell.paragraphs[0].clear()
        p2 = desc_cell.paragraphs[0]
        r2 = p2.add_run(co_text or "")
        r2.font.size = Pt(10)

        _set_cell_border(label_cell)
        _set_cell_border(desc_cell)
