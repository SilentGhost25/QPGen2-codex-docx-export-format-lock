from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = OxmlElement(tag)
        element.set(qn("w:val"), kwargs.get(edge, "single"))
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")
        tcBorders.append(element)
    tcPr.append(tcBorders)


def set_cell_background(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def set_column_widths(table, widths_in_inches):
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths_in_inches):
                cell.width = Inches(widths_in_inches[idx])


def add_paragraph_to_cell(cell, text, bold=False, font_size=11,
                            alignment=WD_ALIGN_PARAGRAPH.LEFT,
                            color=None, italic=False):
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    p.alignment = alignment
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p


def build_header(document: Document, meta: dict):
    """
    Builds the complete static header of the question paper.
    """

    # ── 1. Logo + Institute name row ──────────────────────────────────────
    logo_table = document.add_table(rows=1, cols=3)
    logo_table.style = "Table Grid"

    left_cell = logo_table.rows[0].cells[0]
    mid_cell = logo_table.rows[0].cells[1]
    right_cell = logo_table.rows[0].cells[2]

    # Left logo
    left_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    left_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if meta.get("left_logo_path") and os.path.exists(meta["left_logo_path"]):
        run = left_cell.paragraphs[0].add_run()
        run.add_picture(meta["left_logo_path"], width=Inches(0.9))
    else:
        add_paragraph_to_cell(
            left_cell, "[DSU Logo]",
            alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=9
        )

    # Centre text block
    mid_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    mid_cell.paragraphs[0].clear()

    lines = [
        (meta.get("institute_name", ""), True, 13),
        (meta.get("autonomous", ""), False, 10),
        (meta.get("affiliated", ""), False, 10),
        (meta.get("approved", ""), False, 10),
        (meta.get("accredited", ""), False, 10),
        (meta.get("nba", ""), False, 10),
        (meta.get("nba_programs", ""), False, 10),
    ]

    first = True
    for text, bold, size in lines:
        if not text:
            continue
        if first:
            p = mid_cell.paragraphs[0]
            first = False
        else:
            p = mid_cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)

    # Right logo (IQAC)
    right_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    right_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if meta.get("right_logo_path") and os.path.exists(meta["right_logo_path"]):
        run = right_cell.paragraphs[0].add_run()
        run.add_picture(meta["right_logo_path"], width=Inches(0.9))
    else:
        add_paragraph_to_cell(
            right_cell, "[IQAC Logo]",
            alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=9
        )

    set_column_widths(logo_table, [1.0, 6.0, 1.0])

    for row in logo_table.rows:
        for cell in row.cells:
            set_cell_border(cell)

    document.add_paragraph()

    # ── 2. USN row ────────────────────────────────────────────────────────
    usn_table = document.add_table(rows=1, cols=2)
    usn_table.style = "Table Grid"
    usn_table.rows[0].cells[0].paragraphs[0].clear()
    p = usn_table.rows[0].cells[0].paragraphs[0]
    r = p.add_run(f"USN: {meta.get('usn', '')}")
    r.bold = True
    r.font.size = Pt(11)
    set_column_widths(usn_table, [4.0, 4.0])

    document.add_paragraph()

    # ── 3. Department heading ─────────────────────────────────────────────
    dept_para = document.add_paragraph()
    dept_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dept_run = dept_para.add_run(meta.get("department", ""))
    dept_run.bold = True
    dept_run.font.size = Pt(13)

    # ── 4. Exam title ─────────────────────────────────────────────────────
    exam_para = document.add_paragraph()
    exam_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    exam_run = exam_para.add_run(meta.get("exam_title", ""))
    exam_run.bold = True
    exam_run.font.size = Pt(13)

    # ── 5. Subject details table ──────────────────────────────────────────
    detail_table = document.add_table(rows=4, cols=6)
    detail_table.style = "Table Grid"

    rows_data = [
        [
            ("Subject:", True), (meta.get("subject", ""), False),
            ("Subject Code:", True), (meta.get("subject_code", ""), False),
            ("Semester:", True), (meta.get("semester", ""), False),
        ],
        [
            ("Max. Marks:", True), (meta.get("max_marks", ""), False),
            ("Batch:", True), (meta.get("batch", ""), False),
            ("Duration:", True), (meta.get("duration", ""), False),
        ],
        [
            ("Date of Exam:", True), (meta.get("date", ""), False),
            ("Teaching Dept.:", True), (meta.get("teaching_department", ""), False),
            ("", False), ("", False),
        ],
        [
            ("RBT Levels:", True),
            (meta.get(
                "rbt_legend",
                "L1-Remember, L2-Understand, L3-Apply, "
                "L4-Analyze, L5-Evaluate, L6-Create"
            ), False),
            ("", False), ("", False), ("", False), ("", False),
        ],
    ]

    for r_idx, row_data in enumerate(rows_data):
        row = detail_table.rows[r_idx]
        for c_idx, (text, bold) in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.paragraphs[0].clear()
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.bold = bold
            run.font.size = Pt(10)
            set_cell_border(cell)

    # Merge the RBT row across cols 1-5
    rbt_row = detail_table.rows[3]
    rbt_row.cells[1].merge(rbt_row.cells[5])

    set_column_widths(detail_table, [1.1, 1.6, 1.1, 1.6, 1.1, 0.9])

    document.add_paragraph()

    # ── 6. Instruction ────────────────────────────────────────────────────
    instr_para = document.add_paragraph()
    instr_run = instr_para.add_run(
        f"Instruction: {meta.get('instruction', 'Answer all questions.')}"
    )
    instr_run.bold = True
    instr_run.font.size = Pt(11)
