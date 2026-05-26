from __future__ import annotations

import logging
from pathlib import Path
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches

logger = logging.getLogger(__name__)

def add_image_to_cell(cell, image_path: str | None, width_inches: float = 3.5) -> None:
    if not image_path:
        return
    path = Path(image_path)
    if not path.exists():
        logger.warning(f"Image path {image_path} does not exist. Skipping image injection.")
        return

    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run()
    try:
        run.add_picture(str(path), width=Inches(width_inches))
        logger.info(f"Successfully rendered image: {image_path}")
    except Exception as e:
        logger.error(f"Error rendering image {image_path}: {e}")
