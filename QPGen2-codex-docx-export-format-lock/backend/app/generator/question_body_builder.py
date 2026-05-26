from dataclasses import dataclass, field
from typing import Optional, List
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─────────────────────────────────────────────────────────────────────────────
# Data model
# ─────────────────────────────────────────────────────────────────────────────

@dataclass
class QuestionSlot:
    qno: int                          # e.g. 1, 2, 3 …
    module: int                       # 1-5
    marks: int
    co: str                           # e.g. "CO1"
    rbt: str                          # e.g. "L2"
    text: str
    subpart: Optional[str] = None     # "a", "b", "c" or None
    image_path: Optional[str] = None
    is_or_boundary: bool = False      # True → insert OR row BEFORE this slot

    @property
    def qno_label(self) -> str:
        if self.subpart:
            return f"{self.qno}({self.subpart})"
        return str(self.qno)


# ─────────────────────────────────────────────────────────────────────────────
# Low-level helpers
# ─────────────────────────────────────────────────────────────────────────────

def _set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _set_row_borders(row):
    for cell in row.cells:
        _set_cell_border(cell)


def _set_cell_background(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _write_cell(cell, text: str, bold=False, size=11,
                align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    para = cell.paragraphs[0]
    para.clear()
    para.alignment = align
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)


def _set_col_widths(table, widths):
    """widths: list of Inches values, one per column."""
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths):
                cell.width = widths[i]


# Column widths  (Q No | Question | Marks | COs | RBTL)
COL_WIDTHS = [
    Inches(0.6),   # Q No
    Inches(5.2),   # Questions
    Inches(0.7),   # Marks
    Inches(0.6),   # COs
    Inches(0.7),   # RBTL
]

HEADER_BG = "D9E1F2"   # light blue for table header


# ─────────────────────────────────────────────────────────────────────────────
# Header row
# ─────────────────────────────────────────────────────────────────────────────

def _add_header_row(table):
    header_row = table.rows[0]
    labels = ["Q No", "Questions", "Marks", "COs", "RBTL"]
    for i, label in enumerate(labels):
        cell = header_row.cells[i]
        _write_cell(cell, label, bold=True, size=11,
                    align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_background(cell, HEADER_BG)
    _set_row_borders(header_row)
    _set_col_widths(table, COL_WIDTHS)


# ─────────────────────────────────────────────────────────────────────────────
# OR separator row
# ─────────────────────────────────────────────────────────────────────────────

def _add_or_row(table):
    row = table.add_row()
    # Merge all 5 cells into one
    merged = row.cells[0]
    for i in range(1, 5):
        merged = merged.merge(row.cells[i])
    _write_cell(merged, "OR", bold=True, size=11,
                align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell_background(merged, "F2F2F2")
    _set_row_borders(row)
    _set_col_widths(table, COL_WIDTHS)


# ─────────────────────────────────────────────────────────────────────────────
# Single question row
# ─────────────────────────────────────────────────────────────────────────────

def _add_question_row(table, slot: QuestionSlot):
    row = table.add_row()
    cells = row.cells

    # Q No
    _write_cell(cells[0], slot.qno_label, bold=True, size=11,
                align=WD_ALIGN_PARAGRAPH.CENTER)

    # Question text (and optional image)
    cells[1].vertical_alignment = WD_ALIGN_VERTICAL.TOP
    para = cells[1].paragraphs[0]
    para.clear()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if slot.image_path:
        try:
            import os
            if os.path.exists(slot.image_path):
                run = para.add_run()
                run.add_picture(slot.image_path, width=Inches(4.5))
                cells[1].add_paragraph()   # spacing after image
            else:
                para.add_run(f"\n[Image not found: {slot.image_path}]\n")
        except Exception:
            para.add_run("\n[Image loading error]\n")

    text_para = cells[1].add_paragraph() if slot.image_path else para
    text_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    text_run = text_para.add_run(slot.text)
    text_run.font.size = Pt(11)

    # Marks
    _write_cell(cells[2], str(slot.marks), size=11,
                align=WD_ALIGN_PARAGRAPH.CENTER)

    # CO
    _write_cell(cells[3], slot.co, size=11,
                align=WD_ALIGN_PARAGRAPH.CENTER)

    # RBTL
    _write_cell(cells[4], slot.rbt, size=11,
                align=WD_ALIGN_PARAGRAPH.CENTER)

    _set_row_borders(row)
    _set_col_widths(table, COL_WIDTHS)


# ─────────────────────────────────────────────────────────────────────────────
# Main builder
# ─────────────────────────────────────────────────────────────────────────────

def build_question_body(document: Document, slots: List[QuestionSlot]):
    """
    Creates the question table from scratch.
    Rows are added dynamically — no fixed row count, no old template.
    """

    # Create table with only the header row
    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"

    _add_header_row(table)

    for slot in slots:
        if slot.is_or_boundary:
            _add_or_row(table)
        _add_question_row(table, slot)

    # Final pass: enforce column widths on every row
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(COL_WIDTHS):
                cell.width = COL_WIDTHS[i]
