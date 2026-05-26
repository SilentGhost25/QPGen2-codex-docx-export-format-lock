from __future__ import annotations

import logging
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

logger = logging.getLogger(__name__)

def _style_run(
    run,
    *,
    size: int = 9,
    bold: bool = False,
    italic: bool = False,
    color: RGBColor | None = None,
) -> None:
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color

def _set_table_borders(
    table,
    *,
    top: bool = True,
    left: bool = True,
    bottom: bool = True,
    right: bool = True,
    inside_h: bool = True,
    inside_v: bool = True,
    size: str = "10",
) -> None:
    borders = OxmlElement("w:tblBorders")
    mapping = {
        "top": top,
        "left": left,
        "bottom": bottom,
        "right": right,
        "insideH": inside_h,
        "insideV": inside_v,
    }
    for name, enabled in mapping.items():
        border = OxmlElement(f"w:{name}")
        border.set(qn("w:val"), "single" if enabled else "nil")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")
        borders.append(border)

    table_element = table._tbl
    table_pr = table_element.tblPr
    existing = table_pr.first_child_found_in("w:tblBorders")
    if existing is not None:
        table_pr.remove(existing)
    table_pr.append(borders)

def _set_table_fixed_layout(table) -> None:
    table_pr = table._tbl.tblPr
    existing = table_pr.first_child_found_in("w:tblLayout")
    if existing is not None:
        table_pr.remove(existing)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    table_pr.append(layout)

def _set_cell_margins(
    cell,
    *,
    top: int = 45,
    left: int = 65,
    bottom: int = 45,
    right: int = 65,
) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in {
        "top": top,
        "start": left,
        "bottom": bottom,
        "end": right,
    }.items():
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")

def _set_cell(
    cell,
    text: str,
    *,
    bold: bool = False,
    italic: bool = False,
    size: int = 9,
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
    vertical_align: WD_ALIGN_VERTICAL = WD_ALIGN_VERTICAL.CENTER,
) -> None:
    cell.text = ""
    _set_cell_margins(cell)
    lines = str(text).splitlines() or [""]
    for index, line in enumerate(lines):
        paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        paragraph.alignment = align
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        run = paragraph.add_run(line)
        _style_run(run, size=size, bold=bold, italic=italic)
    cell.vertical_alignment = vertical_align
