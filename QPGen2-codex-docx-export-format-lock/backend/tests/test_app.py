import os
import tempfile
from io import BytesIO
from pathlib import Path

from docx import Document as DocxDocument
from fastapi.testclient import TestClient

os.environ["DATABASE_URL"] = "sqlite://"

test_workspace_root = Path(tempfile.gettempdir()) / "qpgen_backend_test"
os.environ["STORAGE_ROOT"] = str(test_workspace_root / "test-storage")

test_workspace_root.mkdir(parents=True, exist_ok=True)
Path(os.environ["STORAGE_ROOT"]).mkdir(parents=True, exist_ok=True)

from app.main import app

def login(client: TestClient, email: str, password: str) -> str:
    response = client.post("/api/v1/auth/login", json={"email": email, "password": password})
    assert response.status_code == 200
    return response.json()["access_token"]


def _question_table(document: DocxDocument):
    return next(
        table
        for table in document.tables
        if table.rows
        and table.rows[0].cells
        and table.rows[0].cells[0].text.strip().replace("\n", " ") == "Q No"
    )


def test_teacher_to_hod_review_flow() -> None:
    with TestClient(app) as client:
        teacher_token = login(client, "teacher@dsatm.edu", "Teacher@123")

        generate_response = client.post(
            "/api/v1/papers/generate",
            headers={"Authorization": f"Bearer {teacher_token}"},
            json={
                "subject_id": 1,
                "title": "IAT-1 Question Paper",
                "exam_type": "IAT-1",
                "semester": "5",
                "batch": "2022-26",
                "max_marks": 30,
                "duration_minutes": 90,
                "teaching_department": "AIML",
                "prompt": "Generate a balanced paper for CO1-CO3 with a mix of L1-L4.",
                "rbt_levels": ["L1", "L2", "L3", "L4"],
                "module_numbers": [1, 2, 3, 4],
            },
        )
        assert generate_response.status_code == 200
        paper = generate_response.json()

        submit_response = client.post(
            f"/api/v1/papers/{paper['id']}/submit",
            headers={"Authorization": f"Bearer {teacher_token}"},
        )
        assert submit_response.status_code == 200
        assert submit_response.json()["status"] == "pending_review"

        hod_token = login(client, "hod@dsatm.edu", "Hod@123")
        review_response = client.post(
            f"/api/v1/reviews/{paper['id']}/action",
            headers={"Authorization": f"Bearer {hod_token}"},
            json={"decision": "approved", "comments": "Balanced and ready for exam cell."},
        )
        assert review_response.status_code == 200
        assert review_response.json()["status"] == "approved"

        download_response = client.get(
            f"/api/v1/papers/{paper['id']}/download",
            headers={"Authorization": f"Bearer {hod_token}"},
        )
        assert download_response.status_code == 200
        exported = DocxDocument(BytesIO(download_response.content))
        paragraph_text = "\n".join(paragraph.text for paragraph in exported.paragraphs)
        table_text = "\n".join(
            cell.text for table in exported.tables for row in table.rows for cell in row.cells
        )
        assert "Dayananda Sagar Academy of Technology & Management" in table_text
        assert "USN:" in table_text
        assert "Department of Artificial Intelligence and Machine Learning" in paragraph_text or table_text
        assert "COs" in table_text
        assert "RBTL" in table_text


def test_end_sem_download_places_or_between_alternative_questions() -> None:
    with TestClient(app) as client:
        teacher_token = login(client, "teacher@dsatm.edu", "Teacher@123")

        generate_response = client.post(
            "/api/v1/ai/generate-paper",
            headers={"Authorization": f"Bearer {teacher_token}"},
            json={
                "subject_id": 3,
                "title": "End-Sem NLP Paper",
                "exam_type": "End-Sem",
                "semester": "6",
                "batch": "2022-26",
                "max_marks": 100,
                "duration_minutes": 180,
                "exam_date": "2026-04-30",
                "teaching_department": "AIML",
                "prompt": "Generate a module-balanced end-sem paper for NLP.",
                "rbt_levels": ["L1", "L2", "L3", "L4", "L5", "L6"],
                "module_numbers": [1, 2, 3, 4, 5],
                "difficulty_distribution": {"easy": 30, "medium": 40, "hard": 30},
                "co_targets": {"CO1": 20, "CO2": 20, "CO3": 20, "CO4": 20, "CO5": 20},
                "use_notes": False,
                "use_question_bank": False,
                "use_syllabus": False,
                "use_previous_papers": False,
            },
        )
        assert generate_response.status_code == 200
        paper = generate_response.json()
        assert 14 <= len(paper["questions"]) <= 22

        download_response = client.get(
            f"/api/v1/papers/{paper['id']}/download",
            headers={"Authorization": f"Bearer {teacher_token}"},
        )
        assert download_response.status_code == 200
        exported = DocxDocument(BytesIO(download_response.content))
        question_table = _question_table(exported)

        labels = [row.cells[0].text.strip() for row in question_table.rows]
        question_rows = {
            row.cells[0].text.strip(): row.cells[1].text.strip()
            for row in question_table.rows
            if row.cells[0].text.strip() and row.cells[0].text.strip()[0].isdigit()
        }
        assert 14 <= len(question_rows) <= 22
        assert any(question_rows.values())
        assert labels.count("OR") == 5


def test_manual_generation_respects_selected_question_order() -> None:
    with TestClient(app) as client:
        teacher_token = login(client, "teacher@dsatm.edu", "Teacher@123")

        questions_response = client.get(
            "/api/v1/questions?subject_id=1",
            headers={"Authorization": f"Bearer {teacher_token}"},
        )
        assert questions_response.status_code == 200
        sorted_qs = sorted(questions_response.json(), key=lambda x: (x.get("module_number") or 1, x.get("id") or 1))
        manual_ids = [item["id"] for item in sorted_qs[:18]]
        assert len(manual_ids) == 18

        generate_response = client.post(
            "/api/v1/ai/generate-paper",
            headers={"Authorization": f"Bearer {teacher_token}"},
            json={
                "subject_id": 1,
                "title": "Manual ML Paper",
                "exam_type": "IAT-2",
                "semester": "5",
                "batch": "2022-26",
                "max_marks": 50,
                "duration_minutes": 90,
                "teaching_department": "AIML",
                "prompt": "Use the manually curated question list.",
                "rbt_levels": ["L1", "L2", "L3", "L4"],
                "module_numbers": [1, 2, 3, 4, 5],
                "manual_question_ids": manual_ids,
            },
        )
        assert generate_response.status_code == 200
        paper = generate_response.json()
        assert [item["question_id"] for item in paper["questions"]] == manual_ids
        assert paper["questions"][0]["section_label"] == "1(a)"
        assert paper["questions"][1]["section_label"] == "1(b)"
        assert paper["questions"][2]["section_label"] == "1(c)"

        download_response = client.get(
            f"/api/v1/papers/{paper['id']}/download",
            headers={"Authorization": f"Bearer {teacher_token}"},
        )
        assert download_response.status_code == 200
        exported = DocxDocument(BytesIO(download_response.content))
        question_table = _question_table(exported)
        labels = [row.cells[0].text.strip() for row in question_table.rows]
        assert "OR" in labels
        assert labels.count("OR") >= 1


def test_figure_rendering_in_docx() -> None:
    from app.database import SessionLocal
    from app.academic.models import AcademicDocument
    from app.models import Question
    from PIL import Image

    # 1. Create a dummy image
    dummy_img_dir = test_workspace_root / "test-storage" / "multimodal_images"
    dummy_img_dir.mkdir(parents=True, exist_ok=True)
    dummy_img_path = dummy_img_dir / "test_figure.png"
    img = Image.new("RGB", (100, 100), color="red")
    img.save(dummy_img_path)

    # 2. Create a question paper containing a question with the figure tag
    with TestClient(app) as client:
        # Seed academic document structured content inside client lifespan context
        db = SessionLocal()
        try:
            doc = AcademicDocument(
                subject_id=1,
                uploaded_by=2,
                file_name="multimodal_test.pdf",
                file_type="pdf",
                storage_path=str(dummy_img_path),
                processing_status="completed",
                structured_content={
                    "pages": [
                        {
                            "page_number": 1,
                            "blocks": [
                                {
                                    "type": "figure",
                                    "caption": "Transformer Architecture",
                                    "image_path": str(dummy_img_path),
                                    "analysis": {"description": "Encoder-Decoder block"}
                                }
                            ]
                        }
                    ]
                }
            )
            db.add(doc)
            db.commit()
        finally:
            db.close()

        teacher_token = login(client, "teacher@dsatm.edu", "Teacher@123")

        questions_response = client.get(
            "/api/v1/questions?subject_id=1",
            headers={"Authorization": f"Bearer {teacher_token}"},
        )
        assert questions_response.status_code == 200
        q_data = questions_response.json()
        manual_ids = [item["id"] for item in q_data[:20]]

        # Update the first question directly in DB
        db = SessionLocal()
        try:
            db_q = db.query(Question).filter(Question.id == manual_ids[0]).first()
            db_q.text = "Explain the Transformer Architecture block as shown in [FIGURE: Transformer Architecture]."
            db.commit()
        finally:
            db.close()

        # Generate paper using manual questions
        generate_response = client.post(
            "/api/v1/ai/generate-paper",
            headers={"Authorization": f"Bearer {teacher_token}"},
            json={
                "subject_id": 1,
                "title": "Multimodal ML Paper",
                "exam_type": "IAT-2",
                "semester": "5",
                "batch": "2022-26",
                "max_marks": 50,
                "duration_minutes": 90,
                "teaching_department": "AIML",
                "prompt": "Use the manually curated question list.",
                "rbt_levels": ["L1", "L2", "L3", "L4"],
                "module_numbers": [1, 2, 3, 4, 5],
                "manual_question_ids": manual_ids,
            },
        )
        assert generate_response.status_code == 200
        paper = generate_response.json()

        # Download paper
        download_response = client.get(
            f"/api/v1/papers/{paper['id']}/download",
            headers={"Authorization": f"Bearer {teacher_token}"},
        )
        assert download_response.status_code == 200

        # Load and verify DOCX contains image and that tag is removed!
        exported = DocxDocument(BytesIO(download_response.content))

        # Verify tag is removed from the text
        question_table = _question_table(exported)
        text_content = "\n".join(cell.text for row in question_table.rows for cell in row.cells)
        assert "[FIGURE" not in text_content
        assert "Explain the Transformer Architecture block as shown in" in text_content

        # Verify image was rendered/inserted into document
        assert len(exported.inline_shapes) > 0, "No image found in the exported DOCX document!"
        print("SUCCESS: Figure image successfully rendered and embedded into the DOCX question paper!")


def test_module_centric_questions_api() -> None:
    with TestClient(app) as client:
        teacher_token = login(client, "teacher@dsatm.edu", "Teacher@123")
        headers = {"Authorization": f"Bearer {teacher_token}"}
        
        # Call the new module-centric questions library API
        response = client.get("/api/v1/ai/module-questions?subject_id=1", headers=headers)
        assert response.status_code == 200, f"Failed: {response.text}"
        data = response.json()
        
        # Verify response structure (should be a list of modules)
        assert isinstance(data, list)
        assert len(data) > 0
        
        # Check first module structure
        first_mod = data[0]
        assert "module" in first_mod
        assert "questions" in first_mod
        assert isinstance(first_mod["questions"], list)
        
        # Verify candidate question fields
        if len(first_mod["questions"]) > 0:
            q = first_mod["questions"][0]
            assert "id" in q
            assert "text" in q
            assert "marks" in q
            assert "co" in q
            assert "rbt" in q
            assert "difficulty" in q
            assert "topic" in q
            assert "source" in q
            assert "recommendation_score" in q
            assert "recommended" in q
            
        # Test filtering by a specific module
        mod_filtered_response = client.get("/api/v1/ai/module-questions?subject_id=1&module=1", headers=headers)
        assert mod_filtered_response.status_code == 200
        filtered_data = mod_filtered_response.json()
        assert isinstance(filtered_data, dict)
        assert filtered_data["module"] == 1
        assert isinstance(filtered_data["questions"], list)


def test_hierarchical_chunking_and_image_pipeline() -> None:
    from app.academic.chunking import semantic_chunk, count_tokens, clean_ocr, is_high_quality, infer_microchunk_type
    from app.academic.question_allocator import allocate_blueprint
    from app.academic.planning.blueprint_engine import QuestionTask

    # 1. Test Chunker Heading Awareness and Cleanups
    text = (
        "MODULE - 3\n"
        "HEURISTIC SEARCH\n"
        "This is a long and highly informative paragraph about Heuristic Search. Page 123.\n"
        "It explains standard heuristic search, A* algorithm, and PEAS representation.\n"
        "FIGURE 12. Diagram showing search transitions..."
    )
    chunks = semantic_chunk(text, min_tokens=5, max_tokens=100)
    assert len(chunks) > 0
    
    first_chunk = chunks[0]
    assert first_chunk.source_section == "HEURISTIC SEARCH"
    assert "Heuristic Search" in first_chunk.topic_name
    
    # Verify OCR margins are cleaned
    assert "Page 123" not in first_chunk.text
    assert "FIGURE 12" not in first_chunk.text
    
    # Verify microchunk type inference
    assert infer_microchunk_type("Define A* search and heuristics.") == "definition"
    assert infer_microchunk_type("The A* search algorithm steps are: ...") == "algorithm"

    # 2. Test Allocator Image Question Priority Matching
    class MockQuestion:
        def __init__(self, id, module_number, marks, co, bloom, image_path=None):
            self.id = id
            self.module_number = module_number
            self.marks = marks
            self.course_outcome = co
            self.bloom_level = bloom
            self.image_path = image_path
            self.difficulty = "balanced"
            self.tags = []

    pool = [
        MockQuestion(1, 1, 5, "CO1", "L2", image_path=None),
        MockQuestion(2, 1, 5, "CO1", "L2", image_path="/static/images/8queens.png"),
        MockQuestion(3, 1, 5, "CO1", "L2", image_path=None),
    ]

    blueprint = [
        QuestionTask(
            question_number=1,
            subpart="a",
            label="1a",
            module=1,
            co="CO1",
            rbt="L2",
            difficulty="balanced",
            marks=5,
        )
    ]

    # Force images for module 1
    results = allocate_blueprint(
        blueprint=blueprint,
        pool=pool,
        module_image_map={1: True}
    )

    assert len(results) == 1
    # Should prioritize MockQuestion 2 which has an image path!
    assert results[0].question.id == 2
    assert results[0].question.image_path == "/static/images/8queens.png"


def test_co_description_synthesis() -> None:
    from sqlalchemy import select
    from app.database import SessionLocal
    from app.models import Subject
    from app.academic.models import KnowledgeChunk, SubjectSyllabus
    from app.academic.co_description_generator import generate_subject_co_descriptions

    db = SessionLocal()
    try:
        # Create a test subject
        subj = Subject(
            dept_id=1,
            name="Artificial Intelligence and Machine Learning Notes",
            code="AIML-501",
            semester=5,
            max_marks=50
        )
        db.add(subj)
        db.commit()
        db.refresh(subj)

        # Add some chunks with topics to module 1 and 2
        chunk1 = KnowledgeChunk(
            document_id=999,
            subject_id=subj.id,
            chunk_text="Intelligent agents must act rationally in a PEAS environment.",
            chunk_index=0,
            token_count=10,
            module_number=1,
            topic_name="Intelligent Agents",
            approval_status="approved"
        )
        chunk2 = KnowledgeChunk(
            document_id=999,
            subject_id=subj.id,
            chunk_text="A* search uses heuristics to guide the search towards goal.",
            chunk_index=1,
            token_count=10,
            module_number=2,
            topic_name="A* Search",
            approval_status="approved"
        )
        db.add_all([chunk1, chunk2])
        db.commit()

        # Run synthesis
        co_descriptions = generate_subject_co_descriptions(db, subj.id)

        # Verify output
        assert "CO1" in co_descriptions
        assert "CO2" in co_descriptions
        assert "Intelligent Agents" in co_descriptions["CO1"]
        # AI template should be selected because name has "Artificial Intelligence"
        assert "intelligent agent architectures" in co_descriptions["CO1"]
        assert "A* Search" in co_descriptions["CO2"]
        assert "search strategies" in co_descriptions["CO2"]

        # Verify it got saved to database syllabus
        syllabus = db.scalar(select(SubjectSyllabus).where(SubjectSyllabus.subject_id == subj.id))
        assert syllabus is not None
        assert syllabus.co_json == co_descriptions

    finally:
        db.close()



